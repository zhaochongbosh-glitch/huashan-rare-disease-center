# -*- coding: utf-8 -*-
"""Generate structured disease-detail DOCX drafts for the website workflow.

Examples:
  python tools/generate_disease_docx.py "戈谢病" --offline-template
  python tools/generate_disease_docx.py "戈谢病"
  python tools/generate_disease_docx.py "戈谢病" "庞贝病" --import-site

With OPENAI_API_KEY set, the script uses the OpenAI Responses API for text and
for image generation. Without an API key, use --offline-template to create a
review-ready blank scaffold that can still be edited manually and imported later.
"""

from __future__ import annotations

import argparse
import base64
import json
import os
import re
import subprocess
import sys
import textwrap
import urllib.error
import urllib.request
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DISEASES_PATH = ROOT / "data" / "diseases.json"
DEFAULT_LOGO = ROOT / "assets" / "huashan-rare-disease-center-logo.jpg"
DEFAULT_OUT_DIR = ROOT / "generated" / "disease-docx"
API_URL = "https://api.openai.com/v1/responses"

TEXT_MODEL = "gpt-5.5"
IMAGE_MODEL = "gpt-5.5"


@dataclass
class CatalogDisease:
    id: str
    batch: int
    catalog_no: int
    name_cn: str
    name_en: str
    specialty_group: str

    @property
    def directory_batch(self) -> str:
        label = "第一批" if self.batch == 1 else "第二批"
        year = "2018年" if self.batch == 1 else "2023年"
        return f"国家{label}罕见病目录（{year}），目录序号第{self.catalog_no}位。"


def clean_name(text: str) -> str:
    return re.sub(r"[\s,，()（）《》“”\"'：:；;、/\\-]+", "", text).lower()


def slugify(text: str) -> str:
    ascii_seed = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return ascii_seed or "disease-detail"


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', "-", name).strip(". ")
    return cleaned or "疾病详情"


def load_catalog() -> list[CatalogDisease]:
    data = json.loads(DISEASES_PATH.read_text(encoding="utf-8"))
    catalog: list[CatalogDisease] = []
    for item in data.get("items", []):
        catalog.append(
            CatalogDisease(
                id=item["id"],
                batch=int(item["batch"]),
                catalog_no=int(item["catalogNo"]),
                name_cn=item["nameCn"],
                name_en=item.get("nameEn", ""),
                specialty_group=item.get("specialtyGroup", "待中心分组"),
            )
        )
    return catalog


def find_catalog(name: str, catalog: list[CatalogDisease]) -> CatalogDisease:
    normalized = clean_name(name)
    exact = [item for item in catalog if clean_name(item.name_cn) == normalized]
    if exact:
        return exact[0]
    contains = [
        item
        for item in catalog
        if normalized in clean_name(item.name_cn) or clean_name(item.name_cn) in normalized
    ]
    if len(contains) == 1:
        return contains[0]
    english = [item for item in catalog if normalized and normalized == clean_name(item.name_en)]
    if english:
        return english[0]
    candidates = "；".join(item.name_cn for item in contains[:8]) if contains else "无"
    raise ValueError(f"未能唯一匹配目录疾病：{name}。候选：{candidates}")


def call_openai_responses(api_key: str, payload: dict[str, Any]) -> dict[str, Any]:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        API_URL,
        data=data,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"OpenAI API error {exc.code}: {body}") from exc


def extract_response_text(response: dict[str, Any]) -> str:
    if isinstance(response.get("output_text"), str):
        return response["output_text"]
    chunks: list[str] = []
    for output in response.get("output", []):
        for content in output.get("content", []) if isinstance(output, dict) else []:
            if isinstance(content, dict) and content.get("type") in {"output_text", "text"}:
                chunks.append(content.get("text", ""))
    return "\n".join(chunks).strip()


def extract_image_base64(response: dict[str, Any]) -> str | None:
    for output in response.get("output", []):
        if isinstance(output, dict) and output.get("type") == "image_generation_call":
            result = output.get("result")
            if isinstance(result, str):
                return result
        for content in output.get("content", []) if isinstance(output, dict) else []:
            if isinstance(content, dict):
                if content.get("type") == "image_generation_call" and isinstance(content.get("result"), str):
                    return content["result"]
                if isinstance(content.get("image_base64"), str):
                    return content["image_base64"]
    return None


def parse_json_from_text(text: str) -> dict[str, Any]:
    stripped = text.strip()
    stripped = re.sub(r"^```(?:json)?\s*", "", stripped)
    stripped = re.sub(r"\s*```$", "", stripped)
    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", stripped, flags=re.S)
        if match:
            return json.loads(match.group(0))
        raise


def text_prompt(disease: CatalogDisease) -> tuple[str, str]:
    instructions = textwrap.dedent(
        """
        你是复旦大学附属华山医院罕见病中心网站的医学科普初稿助手。
        请生成适用于网页展示和专家审核的罕见病详情草稿。

        重要要求：
        1. 面向患者和家属，语言准确、克制、清晰，不夸大疗效。
        2. 不给出具体处方、剂量、商业药品推荐或替代医生诊疗的建议。
        3. 对治疗、医保、临床试验、MDT 适应证使用“需由专科医生/医院正式信息确认”的谨慎表述。
        4. 输出必须是严格 JSON，不要 Markdown，不要解释。
        5. 字段必须完整；每个数组保留 3-6 个要点，避免长篇堆砌。
        """
    ).strip()

    user = {
        "任务": "按固定格式撰写疾病详情，适用于华山医院罕见病中心网站展示。",
        "疾病名称": disease.name_cn,
        "所属目录批次": disease.directory_batch,
        "英文名": disease.name_en,
        "建议字段": [
            "疾病名称",
            "所属目录批次",
            "英文名",
            "疾病简介",
            "常见症状",
            "疾病特点",
            "建议就诊方向",
            "初诊材料",
            "是否适合 MDT",
            "相关 MDT 团队",
            "疾病分型",
            "可关联临床研究",
            "政策医保提示",
            "患者常见问题（FAQ）",
            "疾病关键词",
        ],
        "JSON结构": {
            "title": "疾病中文名（英文名）",
            "summary": ["2-4段疾病简介"],
            "symptoms": [{"title": "症状类别", "items": ["症状要点"]}],
            "features": ["疾病特点要点"],
            "visit": {"primary": ["首选科室"], "related": ["相关科室"]},
            "materials": [{"title": "材料类别", "items": ["资料清单"]}],
            "mdt": {"suitable": "是否适合 MDT 的说明", "teams": ["相关 MDT 团队"], "notes": ["MDT补充说明"]},
            "classification": [{"title": "分型类别", "items": ["分型要点"]}],
            "research": [{"title": "研究方向", "items": ["可关联临床研究方向"]}],
            "policy": [{"title": "政策医保提示", "items": ["政策和医保提示"]}],
            "faq": [{"question": "患者常见问题", "answer": "简洁回答"}],
            "keywords": ["关键词"],
        },
    }
    return instructions, json.dumps(user, ensure_ascii=False)


def generate_text_with_openai(disease: CatalogDisease, api_key: str, model: str) -> dict[str, Any]:
    instructions, prompt = text_prompt(disease)
    payload = {
        "model": model,
        "instructions": instructions,
        "input": prompt,
        "reasoning": {"effort": "low"},
    }
    response = call_openai_responses(api_key, payload)
    text = extract_response_text(response)
    if not text:
        raise RuntimeError("OpenAI response did not contain text output.")
    return normalize_detail_json(parse_json_from_text(text), disease)


def offline_template(disease: CatalogDisease) -> dict[str, Any]:
    name = disease.name_cn
    return normalize_detail_json(
        {
            "title": f"{name}（{disease.name_en}）" if disease.name_en else name,
            "summary": [
                f"{name}已收录于{disease.directory_batch}本段为自动生成的编辑模板，需由医学专家补充疾病定义、病因机制、诊断路径和治疗原则。",
                "本页内容定位为患者科普和就诊准备说明，不替代医生面对面诊疗。正式发布前需完成中心专家医学审核。",
            ],
            "symptoms": [{"title": "常见症状", "items": ["待补充典型症状", "待补充需要及时就医的表现", "待补充疾病进展相关表现"]}],
            "features": ["待补充遗传学或病理生理特点", "待补充诊断难点", "待补充长期管理重点"],
            "visit": {"primary": ["待中心确认首选科室"], "related": ["医学遗传科", "相关专科门诊"]},
            "materials": [
                {
                    "title": "初诊材料",
                    "items": [
                        "身份证件与既往门诊病史",
                        "出院小结、病理或穿刺报告",
                        "影像资料原片、纸质报告和原始光盘",
                        "按时间顺序整理的化验、免疫和基因检测结果",
                        "当前用药名称、剂量及病情变化说明",
                    ],
                }
            ],
            "mdt": {
                "suitable": "疑难诊断、多系统受累、治疗方案需多学科共同评估或外院诊断后需复核者，可由专科医生评估是否进入 MDT。",
                "teams": ["罕见病 MDT", "相关专科 MDT"],
                "notes": ["具体团队名称、院区和申请路径需以医院正式门诊信息为准。"],
            },
            "classification": [{"title": "疾病分型", "items": ["待补充临床分型", "待补充基因或病理分型"]}],
            "research": [{"title": "可关联临床研究", "items": ["自然史研究", "登记随访研究", "诊疗路径和生活质量研究"]}],
            "policy": [{"title": "政策医保提示", "items": ["是否纳入医保、特殊药品或救助政策需以最新官方信息和医院审核口径为准。"]}],
            "faq": [
                {"question": f"{name}需要做基因检测吗？", "answer": "部分罕见病需要结合基因检测、临床表现和专科评估综合判断，是否检测及检测项目应由医生决定。"},
                {"question": f"{name}适合申请 MDT 吗？", "answer": "如存在诊断不明确、多系统受累或治疗方案复杂，可先挂相关专科，由医生评估是否申请 MDT。"},
            ],
            "keywords": [name, disease.name_en, "罕见病", "MDT", "华山医院罕见病中心"],
        },
        disease,
    )


def normalize_detail_json(data: dict[str, Any], disease: CatalogDisease) -> dict[str, Any]:
    normalized = dict(data)
    normalized["disease_name"] = disease.name_cn
    normalized["directory_batch"] = disease.directory_batch
    normalized["english_name"] = disease.name_en
    normalized.setdefault("title", f"{disease.name_cn}（{disease.name_en}）" if disease.name_en else disease.name_cn)
    for key in ["summary", "features", "keywords"]:
        value = normalized.get(key, [])
        normalized[key] = value if isinstance(value, list) else [str(value)]
    for key in ["symptoms", "materials", "classification", "research", "policy"]:
        normalized[key] = normalize_group_list(normalized.get(key), key)
    normalized["visit"] = normalize_visit(normalized.get("visit"))
    normalized["mdt"] = normalize_mdt(normalized.get("mdt"))
    normalized["faq"] = normalize_faq(normalized.get("faq"))
    return normalized


def normalize_group_list(value: Any, fallback_title: str) -> list[dict[str, list[str]]]:
    if not value:
        return []
    if isinstance(value, list):
        groups = []
        for index, item in enumerate(value, start=1):
            if isinstance(item, dict):
                title = str(item.get("title") or f"{fallback_title}{index}")
                items = item.get("items") or []
                if isinstance(items, str):
                    items = [items]
                groups.append({"title": title, "items": [str(x) for x in items if str(x).strip()]})
            else:
                groups.append({"title": fallback_title, "items": [str(item)]})
        return groups
    return [{"title": fallback_title, "items": [str(value)]}]


def normalize_visit(value: Any) -> dict[str, list[str]]:
    if not isinstance(value, dict):
        return {"primary": [], "related": []}
    return {
        "primary": as_str_list(value.get("primary")),
        "related": as_str_list(value.get("related")),
    }


def normalize_mdt(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {"suitable": str(value or ""), "teams": [], "notes": []}
    return {
        "suitable": str(value.get("suitable") or ""),
        "teams": as_str_list(value.get("teams")),
        "notes": as_str_list(value.get("notes")),
    }


def normalize_faq(value: Any) -> list[dict[str, str]]:
    if not value:
        return []
    faqs = []
    for item in value if isinstance(value, list) else [value]:
        if isinstance(item, dict):
            question = str(item.get("question") or "").strip()
            answer = str(item.get("answer") or "").strip()
            if question:
                faqs.append({"question": question, "answer": answer})
        else:
            faqs.append({"question": str(item), "answer": ""})
    return faqs


def as_str_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    return [str(value)]


def image_prompt(disease: CatalogDisease) -> str:
    return textwrap.dedent(
        f"""
        创建一幅用于医院患者科普网页的医学示意图，主题是“{disease.name_cn} / {disease.name_en}”。
        竖版 3:4 构图，干净、专业、温和，适合罕见病中心网站。
        使用医学信息图风格，包含抽象的人体系统、细胞/基因/器官机制元素和就诊管理意象。
        不要出现具体药品品牌、处方剂量、恐怖画面、过度写实病痛表情。
        不要在图中写大段文字；仅保留简洁视觉符号，后期会叠加中心 logo 水印。
        """
    ).strip()


def generate_image_with_openai(disease: CatalogDisease, api_key: str, model: str) -> Image.Image:
    payload = {
        "model": model,
        "input": image_prompt(disease),
        "tools": [{"type": "image_generation"}],
    }
    response = call_openai_responses(api_key, payload)
    encoded = extract_image_base64(response)
    if not encoded:
        raise RuntimeError("OpenAI response did not contain image output.")
    return Image.open(BytesIO(base64.b64decode(encoded))).convert("RGB")


def placeholder_image(disease: CatalogDisease) -> Image.Image:
    width, height = 900, 1200
    image = Image.new("RGB", (width, height), "#F7FBFC")
    draw = ImageDraw.Draw(image)
    accent = "#B0403A"
    teal = "#3E7D88"
    navy = "#1D3444"
    draw.rounded_rectangle((52, 52, width - 52, height - 52), radius=36, outline="#D6E6EA", width=4)
    draw.ellipse((135, 210, 765, 840), fill="#EAF5F6", outline="#C7DEE3", width=3)
    draw.line((240, 525, 660, 525), fill=teal, width=8)
    draw.line((450, 315, 450, 735), fill=teal, width=8)
    for angle in range(0, 360, 45):
        x = 450 + int(195 * __import__("math").cos(__import__("math").radians(angle)))
        y = 525 + int(195 * __import__("math").sin(__import__("math").radians(angle)))
        draw.ellipse((x - 42, y - 42, x + 42, y + 42), fill="#FFFFFF", outline="#BDD9DF", width=3)
    draw.ellipse((350, 425, 550, 625), fill="#FFFFFF", outline=accent, width=5)
    font_title = load_font(52)
    font_sub = load_font(28)
    draw.text((90, 900), disease.name_cn[:18], fill=navy, font=font_title)
    if disease.name_en:
        wrapped = textwrap.wrap(disease.name_en, width=34)
        for i, line in enumerate(wrapped[:3]):
            draw.text((92, 982 + i * 38), line, fill="#5D7180", font=font_sub)
    draw.text((92, 1100), "AI 示意图草稿 · 待审核", fill=accent, font=font_sub)
    return image


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def fit_to_ratio(image: Image.Image, width: int = 900, height: int = 1200) -> Image.Image:
    image = image.convert("RGB")
    target_ratio = width / height
    source_ratio = image.width / image.height
    if source_ratio > target_ratio:
        new_width = int(image.height * target_ratio)
        left = (image.width - new_width) // 2
        image = image.crop((left, 0, left + new_width, image.height))
    elif source_ratio < target_ratio:
        new_height = int(image.width / target_ratio)
        top = (image.height - new_height) // 2
        image = image.crop((0, top, image.width, top + new_height))
    return image.resize((width, height), Image.Resampling.LANCZOS)


def add_logo_watermark(image: Image.Image, logo_path: Path) -> Image.Image:
    image = fit_to_ratio(image)
    if not logo_path.exists():
        return image
    logo = Image.open(logo_path).convert("RGBA")
    max_width = int(image.width * 0.22)
    scale = max_width / logo.width
    logo = logo.resize((max_width, max(1, int(logo.height * scale))), Image.Resampling.LANCZOS)
    alpha = logo.getchannel("A")
    alpha = alpha.point(lambda value: int(value * 0.72))
    logo.putalpha(alpha)
    margin = 34
    badge = Image.new("RGBA", (logo.width + 24, logo.height + 24), (255, 255, 255, 185))
    badge.alpha_composite(logo, (12, 12))
    image_rgba = image.convert("RGBA")
    image_rgba.alpha_composite(badge, (image.width - badge.width - margin, image.height - badge.height - margin))
    return image_rgba.convert("RGB")


def save_infographic(image: Image.Image, out_dir: Path, disease: CatalogDisease) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = slugify(disease.name_en) or safe_filename(disease.name_cn)
    target = out_dir / f"{stem}-infographic.jpg"
    image.save(target, "JPEG", quality=90, optimize=True)
    return target


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in [
        ("Heading 1", 18, "1D3444"),
        ("Heading 2", 13, "3E7D88"),
        ("Heading 3", 11.5, "1D3444"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(176, 64, 58)


def add_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(str(line))
        p.paragraph_format.space_after = Pt(3)


def add_groups(doc: Document, groups: list[dict[str, list[str]]]) -> None:
    for group in groups:
        title = str(group.get("title") or "").strip()
        if title:
            p = doc.add_paragraph(title)
            p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
        for item in group.get("items", []):
            doc.add_paragraph(str(item), style=None)


def create_docx(disease: CatalogDisease, detail: dict[str, Any], image_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)

    title = detail.get("title") or disease.name_cn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(title))
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(29, 52, 68)

    subtitle = doc.add_paragraph("华山医院罕见病中心疾病知识库草稿 · 待中心专家医学审核")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(93, 113, 128)
    subtitle.runs[0].font.size = Pt(10)

    add_label(doc, "疾病名称")
    add_lines(doc, [disease.name_cn])
    add_label(doc, "所属目录批次")
    add_lines(doc, [disease.directory_batch])
    add_label(doc, "英文名")
    english_lines = [disease.name_en] if disease.name_en else []
    add_lines(doc, english_lines)

    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(4.0))
    caption = doc.add_paragraph(f"{disease.name_cn}医学科普示意图（AI 草稿，正式发布前需审核）")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(93, 113, 128)

    add_label(doc, "疾病简介")
    add_lines(doc, as_str_list(detail.get("summary")))
    add_label(doc, "常见症状")
    add_groups(doc, detail.get("symptoms", []))
    add_label(doc, "疾病特点")
    add_lines(doc, as_str_list(detail.get("features")))
    add_label(doc, "建议就诊方向")
    add_lines(doc, ["首选科室"])
    add_lines(doc, detail.get("visit", {}).get("primary", []))
    add_lines(doc, ["相关科室"])
    add_lines(doc, detail.get("visit", {}).get("related", []))
    add_label(doc, "初诊材料")
    add_groups(doc, detail.get("materials", []))
    add_label(doc, "是否适合 MDT")
    mdt = detail.get("mdt", {})
    add_lines(doc, [mdt.get("suitable", "")])
    add_lines(doc, ["相关 MDT 团队"])
    add_lines(doc, mdt.get("teams", []))
    add_lines(doc, mdt.get("notes", []))
    add_label(doc, "疾病分型")
    add_groups(doc, detail.get("classification", []))
    add_label(doc, "可关联临床研究")
    add_groups(doc, detail.get("research", []))
    add_label(doc, "政策医保提示")
    add_groups(doc, detail.get("policy", []))
    add_label(doc, "患者常见问题（FAQ）")
    for index, faq in enumerate(detail.get("faq", []), start=1):
        add_lines(doc, [f"{index}. {faq.get('question', '').strip()}"])
        add_lines(doc, [faq.get("answer", "").strip()])
    add_label(doc, "疾病关键词")
    add_lines(doc, ["｜".join(as_str_list(detail.get("keywords")))])

    target = out_dir / f"{safe_filename(disease.name_cn)}.docx"
    doc.save(target)
    return target


def generate_one(
    disease_name: str,
    catalog: list[CatalogDisease],
    out_dir: Path,
    logo_path: Path,
    api_key: str | None,
    text_model: str,
    image_model: str,
    offline_template_mode: bool,
    skip_image_api: bool,
) -> Path:
    disease = find_catalog(disease_name, catalog)
    disease_dir = out_dir / safe_filename(disease.name_cn)
    disease_dir.mkdir(parents=True, exist_ok=True)

    if api_key and not offline_template_mode:
        print(f"text=OpenAI {text_model} | disease={disease.name_cn}")
        detail = generate_text_with_openai(disease, api_key, text_model)
    else:
        print(f"text=offline-template | disease={disease.name_cn}")
        detail = offline_template(disease)

    image: Image.Image
    if api_key and not offline_template_mode and not skip_image_api:
        print(f"image=OpenAI {image_model} | disease={disease.name_cn}")
        image = generate_image_with_openai(disease, api_key, image_model)
    else:
        print(f"image=placeholder | disease={disease.name_cn}")
        image = placeholder_image(disease)
    image = add_logo_watermark(image, logo_path)
    image_path = save_infographic(image, disease_dir, disease)
    docx_path = create_docx(disease, detail, image_path, disease_dir)

    manifest = {
        "diseaseId": disease.id,
        "diseaseName": disease.name_cn,
        "directoryBatch": disease.directory_batch,
        "docx": str(docx_path.relative_to(ROOT)).replace("\\", "/"),
        "image": str(image_path.relative_to(ROOT)).replace("\\", "/"),
        "textMode": "openai" if api_key and not offline_template_mode else "offline-template",
        "imageMode": "openai" if api_key and not offline_template_mode and not skip_image_api else "placeholder",
        "reviewStatus": "AI生成草稿，待中心专家医学审核",
    }
    (disease_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return docx_path


def import_site(docx_paths: list[Path], python_exe: str) -> None:
    importer = ROOT / "tools" / "import_disease_docx.py"
    for path in docx_paths:
        subprocess.run([python_exe, str(importer), str(path), "--dry-run"], cwd=ROOT, check=True)
        subprocess.run([python_exe, str(importer), str(path)], cwd=ROOT, check=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate disease-detail DOCX drafts from disease names.")
    parser.add_argument("diseases", nargs="+", help="Disease names to generate, e.g. 戈谢病")
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR, help="Output directory for generated DOCX files")
    parser.add_argument("--logo", type=Path, default=DEFAULT_LOGO, help="Logo image used as watermark")
    parser.add_argument("--text-model", default=TEXT_MODEL, help="OpenAI model for text generation")
    parser.add_argument("--image-model", default=IMAGE_MODEL, help="OpenAI model for image generation via image_generation tool")
    parser.add_argument("--offline-template", action="store_true", help="Create editable scaffold without OpenAI calls")
    parser.add_argument("--skip-image-api", action="store_true", help="Use placeholder image even when OPENAI_API_KEY is set")
    parser.add_argument("--import-site", action="store_true", help="Run tools/import_disease_docx.py after generation")
    parser.add_argument("--python", default=sys.executable, help="Python executable used for --import-site")
    args = parser.parse_args()

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key and not args.offline_template:
        print("warning=OPENAI_API_KEY is not set; falling back to --offline-template behavior.")

    catalog = load_catalog()
    generated: list[Path] = []
    for disease_name in args.diseases:
        generated.append(
            generate_one(
                disease_name=disease_name,
                catalog=catalog,
                out_dir=args.out_dir.resolve(),
                logo_path=args.logo.resolve(),
                api_key=api_key,
                text_model=args.text_model,
                image_model=args.image_model,
                offline_template_mode=args.offline_template or not bool(api_key),
                skip_image_api=args.skip_image_api,
            )
        )

    if args.import_site:
        import_site(generated, args.python)

    for path in generated:
        print(f"docx={path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
