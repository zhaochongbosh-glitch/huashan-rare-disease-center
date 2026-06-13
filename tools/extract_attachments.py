import json
import shutil
from pathlib import Path

import docx
import openpyxl


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\zhaochob\Downloads\罕见病网站门诊提供材料")

DOCX_PATH = SOURCE / "附件-1 全院区（罕见病）MDT目录.docx"
XLSX_PATH = SOURCE / "附件-2 华山医院未成年患者资质与业务范围0305.xlsx"
FLOORPLAN_DIR = SOURCE / "附件-3-华山医院各院区门诊平面图"


def non_empty_text(value):
    return "" if value is None else str(value).strip()


def extract_docx_tables():
    document = docx.Document(DOCX_PATH)
    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            values = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            if any(values):
                rows.append(values)
        tables.append(
            {
                "tableIndex": table_index,
                "rowCount": len(rows),
                "columnCount": len(table.columns),
                "rows": rows,
            }
        )

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return {"paragraphs": paragraphs, "tables": tables}


def extract_xlsx_sheets():
    workbook = openpyxl.load_workbook(XLSX_PATH, read_only=True, data_only=True)
    sheets = []
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = [non_empty_text(value) for value in row]
            if any(values):
                rows.append(values)
        sheets.append(
            {
                "name": worksheet.title,
                "rowCount": len(rows),
                "columnCount": worksheet.max_column,
                "rows": rows,
            }
        )
    return {"sheets": sheets}


def copy_floorplans():
    output_dir = ROOT / "assets" / "floorplans"
    output_dir.mkdir(parents=True, exist_ok=True)
    copied = []
    for source_file in FLOORPLAN_DIR.rglob("*"):
        if not source_file.is_file():
            continue
        if source_file.suffix.lower() not in {".jpg", ".jpeg", ".png", ".webp"}:
            continue
        campus = source_file.parent.name
        campus_dir = output_dir / campus
        campus_dir.mkdir(parents=True, exist_ok=True)
        safe_name = source_file.name.replace(" ", "-")
        target = campus_dir / safe_name
        shutil.copy2(source_file, target)
        copied.append(
            {
                "campus": campus,
                "title": source_file.stem,
                "file": str(target.relative_to(ROOT)).replace("\\", "/"),
                "sizeBytes": target.stat().st_size,
            }
        )
    copied.sort(key=lambda item: (item["campus"], item["title"]))
    return {"items": copied}


def build_mdt_directory(raw):
    rows = raw["tables"][0]["rows"]
    campuses = rows[1]
    items_by_campus = {}
    for group_start in range(0, 12, 3):
      campus = campuses[group_start].replace("华山总院", "总院")
      items_by_campus.setdefault(campus, [])
      for row in rows[4:]:
          number, name, online = row[group_start:group_start + 3]
          if not number and not name:
              continue
          items_by_campus[campus].append(
              {
                  "no": number,
                  "name": name,
                  "onlineBooking": online == "✓",
                  "onlineBookingLabel": "支持线上预约" if online == "✓" else "暂未开放线上预约",
              }
          )
    return {
        "metadata": {
            "sourceFile": str(DOCX_PATH),
            "updatedAt": "2026-06-12",
            "notes": "由附件 1 提取。线上预约字段按原表“✓/无”转换。",
        },
        "campuses": [
            {"name": campus, "teamCount": len(teams), "teams": teams}
            for campus, teams in items_by_campus.items()
        ],
    }


def build_pediatric_scope(raw):
    sheets = {sheet["name"]: sheet for sheet in raw["sheets"]}
    detail_rows = sheets["Sheet2"]["rows"]
    records = []
    current_category = ""
    current_department = ""
    campus_columns = {
        "总院": {"qualification": 3, "outpatient": 4, "emergency": 5, "inpatient": 6},
        "虹桥院区": {"qualification": 7, "outpatient": 8, "emergency": 9, "inpatient": 10},
        "浦东院区": {"qualification": 11, "outpatient": 12, "inpatient": 13},
        "江苏路分部": {"qualification": 14, "outpatient": 15, "inpatient": 16},
    }
    for row in detail_rows[4:]:
        if row[0]:
            current_category = row[0]
        if row[1]:
            current_department = row[1]
        subject = row[2]
        if not subject:
            continue
        campuses = {}
        for campus, cols in campus_columns.items():
            campuses[campus] = {
                key: row[index] if index < len(row) else ""
                for key, index in cols.items()
            }
        records.append(
            {
                "category": current_category,
                "department": current_department,
                "pediatricSubject": subject,
                "campuses": campuses,
            }
        )

    special_settings = []
    for row in sheets["Sheet3"]["rows"][1:]:
        special_settings.append(
            {
                "clinicOrDoctor": row[0],
                "requestedSetting": row[1] if len(row) > 1 else "",
                "publishedSetting": row[2] if len(row) > 2 else "",
            }
        )

    return {
        "metadata": {
            "sourceFile": str(XLSX_PATH),
            "updatedAt": "2026-06-12",
            "notes": "由附件 2 的 Sheet2、Sheet3 提取。年龄字段沿用原表口径，未开展表示该院区暂未开展对应业务。",
        },
        "records": records,
        "specialSettings": special_settings,
    }


def write_json(relative_path, payload):
    target = ROOT / relative_path
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    mdt = extract_docx_tables()
    pediatric = extract_xlsx_sheets()
    floorplans = copy_floorplans()
    mdt_directory = build_mdt_directory(mdt)
    pediatric_scope = build_pediatric_scope(pediatric)

    write_json(
        "data/mdt-directory.raw.json",
        {
            "metadata": {
                "sourceFile": str(DOCX_PATH),
                "notes": "Raw extraction from attachment 1. Review structure before using as public-facing data.",
            },
            **mdt,
        },
    )
    write_json(
        "data/pediatric-scope.raw.json",
        {
            "metadata": {
                "sourceFile": str(XLSX_PATH),
                "notes": "Raw extraction from attachment 2. Review structure before using as public-facing data.",
            },
            **pediatric,
        },
    )
    write_json(
        "data/floorplans.json",
        {
            "metadata": {
                "sourceDirectory": str(FLOORPLAN_DIR),
                "notes": "Copied from attachment 3 into project assets.",
            },
            **floorplans,
        },
    )
    write_json("data/mdt-directory.json", mdt_directory)
    write_json("data/pediatric-scope.json", pediatric_scope)

    print(f"mdtTables={len(mdt['tables'])}")
    print(f"mdtTeams={sum(c['teamCount'] for c in mdt_directory['campuses'])}")
    print(f"pediatricSheets={len(pediatric['sheets'])}")
    print(f"pediatricRecords={len(pediatric_scope['records'])}")
    print(f"floorplans={len(floorplans['items'])}")


if __name__ == "__main__":
    main()
